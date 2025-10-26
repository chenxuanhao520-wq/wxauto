"""
增强文档处理器
支持多种文档格式解析，使用pandas和各种专业Parser库
"""
import logging
import asyncio
from typing import Dict, List, Any, Optional
from pathlib import Path
from dataclasses import dataclass
import hashlib
from datetime import datetime

# 标准库
import re
import json

# 尝试导入可选依赖
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    pd = None

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    DocxDocument = None

try:
    import PyPDF2
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None
    pdf_extract_text = None

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    openpyxl = None

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    BeautifulSoup = None

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档"""
    document_id: str
    title: str
    content: str
    metadata: Dict[str, Any]
    chunks: List[Dict[str, Any]]
    format: str
    parser_used: str
    parsed_at: datetime


class DocumentProcessor:
    """
    增强文档处理器
    
    功能：
    1. 支持多种文档格式（PDF, DOCX, XLSX, HTML, MD, TXT）
    2. 使用专业Parser库
    3. 智能分块策略
    4. 元数据提取
    5. 表格数据处理（pandas）
    """
    
    def __init__(self):
        """初始化文档处理器"""
        self.supported_formats = self._detect_supported_formats()
        
        logger.info(
            f"文档处理器初始化完成: "
            f"支持格式 {', '.join(self.supported_formats)}"
        )
    
    def _detect_supported_formats(self) -> List[str]:
        """检测支持的格式"""
        formats = ['txt', 'md']  # 基础格式
        
        if PDF_AVAILABLE:
            formats.extend(['pdf'])
        if PYTHON_DOCX_AVAILABLE:
            formats.extend(['docx', 'doc'])
        if EXCEL_AVAILABLE or PANDAS_AVAILABLE:
            formats.extend(['xlsx', 'xls', 'csv'])
        if HTML_AVAILABLE:
            formats.extend(['html', 'htm'])
        
        return formats
    
    async def process_file(
        self,
        file_path: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> ParsedDocument:
        """
        处理文档文件
        
        Args:
            file_path: 文件路径
            chunk_size: 分块大小
            chunk_overlap: 分块重叠
        
        Returns:
            解析后的文档对象
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件扩展名
        ext = path.suffix.lower().lstrip('.')
        
        if ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        # 根据格式选择解析器
        if ext == 'pdf':
            return await self._parse_pdf(file_path, chunk_size, chunk_overlap)
        elif ext in ['docx', 'doc']:
            return await self._parse_docx(file_path, chunk_size, chunk_overlap)
        elif ext in ['xlsx', 'xls', 'csv']:
            return await self._parse_excel(file_path, chunk_size, chunk_overlap)
        elif ext in ['html', 'htm']:
            return await self._parse_html(file_path, chunk_size, chunk_overlap)
        elif ext == 'md':
            return await self._parse_markdown(file_path, chunk_size, chunk_overlap)
        else:  # txt
            return await self._parse_text(file_path, chunk_size, chunk_overlap)
    
    async def _parse_pdf(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> ParsedDocument:
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF解析库未安装，请安装: pip install PyPDF2 pdfminer.six")
        
        path = Path(file_path)
        
        try:
            # 使用pdfminer提取文本（更准确）
            text = pdf_extract_text(file_path)
            
            # 提取元数据
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
                }
            
            # 分块
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            
            # 生成文档ID
            document_id = self._generate_document_id(file_path)
            
            return ParsedDocument(
                document_id=document_id,
                title=path.stem,
                content=text,
                metadata=metadata,
                chunks=chunks,
                format='pdf',
                parser_used='pdfminer',
                parsed_at=datetime.now()
            )
            
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise
    
    async def _parse_docx(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> ParsedDocument:
        """解析DOCX文件"""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("DOCX解析库未安装，请安装: pip install python-docx")
        
        path = Path(file_path)
        
        try:
            doc = DocxDocument(file_path)
            
            # 提取文本
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n\n'.join(paragraphs)
            
            # 提取元数据
            metadata = {
                'paragraph_count': len(paragraphs),
                'core_properties': {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else ''
                }
            }
            
            # 分块
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            
            # 生成文档ID
            document_id = self._generate_document_id(file_path)
            
            return ParsedDocument(
                document_id=document_id,
                title=path.stem,
                content=text,
                metadata=metadata,
                chunks=chunks,
                format='docx',
                parser_used='python-docx',
                parsed_at=datetime.now()
            )
            
        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            raise
    
    async def _parse_excel(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> ParsedDocument:
        """解析Excel文件（使用pandas）"""
        if not PANDAS_AVAILABLE:
            raise RuntimeError("pandas未安装，请安装: pip install pandas openpyxl")
        
        path = Path(file_path)
        
        try:
            # 使用pandas读取
            if path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path, sheet_name=None)  # 读取所有sheet
            
            # 转换为文本
            if isinstance(df, dict):  # 多个sheet
                text_parts = []
                for sheet_name, sheet_df in df.items():
                    text_parts.append(f"## {sheet_name}\n")
                    text_parts.append(sheet_df.to_string())
                text = '\n\n'.join(text_parts)
                
                metadata = {
                    'sheet_count': len(df),
                    'sheet_names': list(df.keys()),
                    'total_rows': sum(len(sheet_df) for sheet_df in df.values())
                }
            else:  # 单个sheet
                text = df.to_string()
                metadata = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': list(df.columns)
                }
            
            # 分块
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            
            # 生成文档ID
            document_id = self._generate_document_id(file_path)
            
            return ParsedDocument(
                document_id=document_id,
                title=path.stem,
                content=text,
                metadata=metadata,
                chunks=chunks,
                format=path.suffix.lstrip('.'),
                parser_used='pandas',
                parsed_at=datetime.now()
            )
            
        except Exception as e:
            logger.error(f"Excel解析失败: {e}")
            raise
    
    async def _parse_html(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> ParsedDocument:
        """解析HTML文件"""
        if not HTML_AVAILABLE:
            raise RuntimeError("HTML解析库未安装，请安装: pip install beautifulsoup4 lxml")
        
        path = Path(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'lxml')
            
            # 提取文本
            text = soup.get_text(separator='\n', strip=True)
            
            # 提取元数据
            title_tag = soup.find('title')
            metadata = {
                'title': title_tag.text if title_tag else '',
                'headings': [h.text for h in soup.find_all(['h1', 'h2', 'h3'])]
            }
            
            # 分块
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            
            # 生成文档ID
            document_id = self._generate_document_id(file_path)
            
            return ParsedDocument(
                document_id=document_id,
                title=path.stem,
                content=text,
                metadata=metadata,
                chunks=chunks,
                format='html',
                parser_used='beautifulsoup4',
                parsed_at=datetime.now()
            )
            
        except Exception as e:
            logger.error(f"HTML解析失败: {e}")
            raise
    
    async def _parse_markdown(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> ParsedDocument:
        """解析Markdown文件"""
        path = Path(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # 提取标题
            headings = re.findall(r'^#+\s+(.+)$', text, re.MULTILINE)
            
            metadata = {
                'headings': headings,
                'heading_count': len(headings)
            }
            
            # 分块
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            
            # 生成文档ID
            document_id = self._generate_document_id(file_path)
            
            return ParsedDocument(
                document_id=document_id,
                title=path.stem,
                content=text,
                metadata=metadata,
                chunks=chunks,
                format='markdown',
                parser_used='builtin',
                parsed_at=datetime.now()
            )
            
        except Exception as e:
            logger.error(f"Markdown解析失败: {e}")
            raise
    
    async def _parse_text(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> ParsedDocument:
        """解析纯文本文件"""
        path = Path(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                'line_count': text.count('\n') + 1,
                'char_count': len(text)
            }
            
            # 分块
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            
            # 生成文档ID
            document_id = self._generate_document_id(file_path)
            
            return ParsedDocument(
                document_id=document_id,
                title=path.stem,
                content=text,
                metadata=metadata,
                chunks=chunks,
                format='text',
                parser_used='builtin',
                parsed_at=datetime.now()
            )
            
        except Exception as e:
            logger.error(f"文本解析失败: {e}")
            raise
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[Dict[str, Any]]:
        """智能分块"""
        # 按段落分割
        paragraphs = re.split(r'\n\s*\n', text)
        
        chunks = []
        current_chunk = ""
        chunk_id = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果当前chunk加上新段落会超过大小
            if len(current_chunk) + len(para) > chunk_size:
                # 保存当前chunk
                if current_chunk:
                    chunks.append({
                        'chunk_id': f'chunk_{chunk_id}',
                        'content': current_chunk.strip(),
                        'char_count': len(current_chunk),
                        'position': chunk_id
                    })
                    chunk_id += 1
                
                # 开始新chunk（保留重叠）
                if chunk_overlap > 0 and current_chunk:
                    overlap_text = current_chunk[-chunk_overlap:]
                    current_chunk = overlap_text + '\n\n' + para
                else:
                    current_chunk = para
            else:
                # 累积段落
                if current_chunk:
                    current_chunk += '\n\n' + para
                else:
                    current_chunk = para
        
        # 保存最后一个chunk
        if current_chunk:
            chunks.append({
                'chunk_id': f'chunk_{chunk_id}',
                'content': current_chunk.strip(),
                'char_count': len(current_chunk),
                'position': chunk_id
            })
        
        return chunks
    
    def _generate_document_id(self, file_path: str) -> str:
        """生成文档ID"""
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
        return f"doc_{file_hash[:12]}"
