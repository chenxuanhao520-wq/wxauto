"""
DOC/DOCX 解析器
"""
import logging
from typing import Dict, Any, List

from .base_parser import BaseParser

logger = logging.getLogger(__name__)


class DocParser(BaseParser):
    """DOC/DOCX 解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析 DOC/DOCX 文件"""
        if not self.validate_file(file_path):
            raise ValueError(f"文件不存在或格式不支持: {file_path}")
        
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")
        
        doc = Document(file_path)
        
        # 提取文本
        text_parts = []
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 识别标题（根据样式）
                if para.style.name.startswith('Heading'):
                    sections.append({
                        'title': text,
                        'level': para.style.name
                    })
                    text_parts.append(f"\n## {text}\n")
                else:
                    text_parts.append(text)
        
        # 提取表格
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(f"\n{table_text}\n")
        
        # 元数据
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'sections': sections
        }
        
        # 添加文档属性（如果有）
        core_props = doc.core_properties
        if core_props.title:
            metadata['title'] = core_props.title
        if core_props.author:
            metadata['author'] = core_props.author
        
        text = "\n".join(text_parts)
        
        logger.info(f"DOC解析成功: {file_path}, {len(text)}字符")
        
        return {
            'text': text,
            'metadata': metadata,
            'sections': sections
        }
    
    def _extract_table(self, table) -> str:
        """提取表格内容"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def supported_formats(self) -> List[str]:
        """支持的格式"""
        return ['.doc', '.docx']

